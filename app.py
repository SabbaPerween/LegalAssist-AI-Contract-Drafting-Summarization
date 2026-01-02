import streamlit as st
import os
from dotenv import load_dotenv
import PyPDF2
import docx
import io
import markdown
import re
from streamlit_quill import st_quill
import time
import html2text
from operator import itemgetter
from fpdf import FPDF 
import fitz
import pytesseract
from PIL import Image  # Add this import for image processing
import shutil 

# --- New Imports for Chatbot Feature ---
# from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import MessagesPlaceholder
from langchain_core.messages import AIMessage, HumanMessage
from langchain.chains.summarize import load_summarize_chain

# --- Tier 1 Imports (Database and Vector Store) ---
from database import *
from langchain_community.vectorstores import FAISS
from langchain_core.runnables import RunnablePassthrough

# --- LangChain & Groq Imports ---
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_huggingface import HuggingFaceEmbeddings

TESSERACT_INSTALLED = shutil.which("tesseract") is not None
# --- Configuration & Initialization ---
load_dotenv()
st.set_page_config(page_title="AI Legal Assistant", layout="wide")
init_db()  # Initialize the SQLite database and tables

# --- RAG & LLM Setup ---
@st.cache_resource
def load_models():
    """Loads the RAG retriever and the Groq LLM. This function is cached for performance."""
    retriever = None
    llm = None
    
    # Load Retriever
    try:
        embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        vector_store = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
        retriever = vector_store.as_retriever()
        st.sidebar.success("Clause library (RAG) loaded!")
    except Exception as e:
        st.sidebar.error("Failed to load RAG index. Please run 'create_vectorstore.py' first.")
        st.sidebar.code(f"Error: {e}")

    # Load LLM
    groq_api_key = os.getenv("GROQ_API_KEY")
    if not groq_api_key:
        st.error("GROQ_API_KEY not found! Please add it to your .env file.")
    else:
        try:
            llm = ChatGroq(temperature=0.7, model_name="llama-3.1-8b-instant", groq_api_key=groq_api_key)
        except Exception as e:
            st.error(f"Failed to configure Groq model. Check API key.")
            st.error(f"Error details: {e}")
    
    return retriever, llm

retriever, llm = load_models()
# app.py

def render_home_page():
    st.title("Welcome to the AI Legal Assistant")
    st.markdown("### Your intelligent partner for drafting, reviewing, and understanding legal documents.")
    
    st.markdown("---")
    
    st.subheader("What would you like to do?")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        with st.container(border=True):
            st.markdown("#### 📄 My Documents")
            st.markdown("View, edit, or delete your previously saved legal documents.")
            if st.button("Go to My Documents", use_container_width=True, key="home_to_docs"):
                set_page("My Documents")
                st.rerun()

    with col2:
        with st.container(border=True):
            st.markdown("#### ✍️ Draft a Contract")
            st.markdown("Generate a new legal document from scratch using AI assistance.")
            if st.button("Start Drafting", use_container_width=True, key="home_to_draft"):
                set_page("Draft")
                st.rerun()

    with col3:
        with st.container(border=True):
            st.markdown("#### 🔍 Review a Document")
            st.markdown("Upload an existing document to chat with it or get a summary.")
            if st.button("Start Reviewing", use_container_width=True, key="home_to_review"):
                set_page("Review")
                st.rerun()

    st.info("💡 **Tip:** You can always use the navigation sidebar on the left to switch between tools.")
    
# In app.py, add this to the HELPER FUNCTIONS section

# In app.py, add to HELPER FUNCTIONS section
def suggest_clauses_for_type(contract_type: str):
    """Makes a quick LLM call to suggest relevant clauses for a given contract type."""
    if not llm: return []
    prompt = ChatPromptTemplate.from_template(
        "List the top 5-7 most critical legal clauses for a '{contract_type}' agreement. "
        "Provide only the clause titles, separated by commas. Do not add any other text or numbering."
    )
    chain = prompt | llm | StrOutputParser()
    try:
        response = chain.invoke({"contract_type": contract_type})
        clauses = [clause.strip() for clause in response.split(',') if clause.strip()]
        return clauses
    except Exception as e:
        print(f"Could not fetch suggested clauses: {e}")
        return []
    
def draft_contract_with_rag(details: dict, persona: str, retriever_instance):
    persona_instructions = {
        "Balanced & Fair": "Adopt a neutral and fair tone, balancing the interests of both parties.",
        "Pro-Party 1": f"Adopt a sharp lawyer persona for {details.get('party1_name', 'Party 1')}. Draft every clause to maximize their advantage and protection.",
        "Pro-Party 2": f"Adopt a sharp lawyer persona for {details.get('party2_name', 'Party 2')}. Draft every clause to maximize their advantage and protection.",
        "Simple English": "Adopt a persona focused on clarity and simplicity. Avoid complex legal jargon where possible, while still creating a legally sound document."
    }
    
    selected_persona_instruction = persona_instructions.get(persona, persona_instructions["Balanced & Fair"])
    user_prompt_parts = ["Draft a formal contract based on these details:"]
    for key, value in details.items():
        if value: user_prompt_parts.append(f"- {key.replace('_', ' ').title()}: {value}")
    user_prompt = "\n".join(user_prompt_parts)

    template = f"""
You are an expert AI legal assistant.
**DRAFTING STYLE:** {selected_persona_instruction}
Your task is to draft a professional legal document.

**CRITICAL INSTRUCTION FOR THE INTRODUCTION PARAGRAPH:**
You MUST construct the opening paragraph of the agreement using the following logic:
- Start with "This {{type}} ('Agreement') is entered into as of {{effective_date}}, by and between".
- For Party 1, state their name: `{{party1_name}}`.
- IF `party1_incorporation_state` is provided, add `, a {{party1_incorporation_state}} entity`.
- IF `party1_address` is provided, add ` with its principal place of business at {{party1_address}}`.
- Conclude with `(hereinafter referred to as "Party 1")`.
- Do the exact same for Party 2.
- If a detail is missing (e.g., incorporation state for an individual), DO NOT mention it or add a placeholder. The sentence must flow naturally.

**ABSOLUTE FORMATTING RULES:**
1.  **Headings:** Use `## Markdown` for all numbered section headings.
2.  **Signature Block:** Use the required two-column `<table>` structure. The content inside the `<td>` tags MUST be plain text with `<br>` tags for line breaks. DO NOT use `<p>` or `<b>` tags inside the `<td>`.

**EXAMPLE OF THE REQUIRED, FPDF-COMPATIBLE SIGNATURE FORMAT:**
<table>
    <tr>
    <td width="50%">
        Party 1:
        <br><br>
        _________________________<br>
        [Signature]<br>
        {{party1_name}}
    </td>
    <td width="50%">
        Party 2:
        <br><br>
        _________________________<br>
        [Signature]<br>
        {{party2_name}}
    </td>
    </tr>
</table>
--------------------------------------------------

**RAG CONTEXT (Use for legal substance, NOT formatting):**
{{context}}

**USER'S DETAILED REQUEST:**
{{user_prompt}}

**--- FINAL CHECK ---**
Generate the complete contract. Ensure the introduction paragraph is perfectly constructed based on the CRITICAL INSTRUCTION, and all formatting rules are followed, especially the ultra-simple, FPDF-compatible signature table.
"""
    prompt = ChatPromptTemplate.from_template(template)
    
    chain_input = {
        "context": itemgetter("user_prompt") | retriever_instance,
        "user_prompt": itemgetter("user_prompt"),
        "type": itemgetter("type"),
        "effective_date": itemgetter("effective_date"),
        "party1_name": itemgetter("party1_name"),
        "party1_address": itemgetter("party1_address"),
        "party1_incorporation_state": itemgetter("party1_incorporation_state"),
        "party2_name": itemgetter("party2_name"),
    }
    
    rag_chain = (chain_input | prompt | llm | StrOutputParser())

    stream_input_dict = details.copy()
    stream_input_dict["user_prompt"] = user_prompt

    return rag_chain.stream(stream_input_dict)


def analyze_contract(contract_text: str, analysis_type: str):
    """
    Analyzes a long contract using the RELIABLE "refine" method. This is slower
    but avoids rate-limiting errors on the free tier by processing chunks sequentially.
    """
    if not llm:
        return "LLM not available. Please check your API key."

    # Split the document into chunks
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=5000, chunk_overlap=200)
    docs = text_splitter.create_documents([contract_text])

    # Define prompts for the "refine" process (initial and refine steps)
    prompts = {
        'Executive Summary': {
            "initial": "Provide a concise summary of the key points, parties, and obligations in the following section of a legal contract:\n\n`{text}`\n\nCONCISE SUMMARY:",
            "refine": "You are a legal analyst. Here is an existing summary: `{existing_answer}`. We have a new section of the contract to consider: `{text}`. Refine the original summary by adding key details from the new section. The goal is a single, cohesive executive summary for the entire document. If the new section is irrelevant, return the original summary.\n\nREFINED SUMMARY:"
        },
        'Clause Breakdown': {
            "initial": "Identify and briefly explain the purpose of any major legal clauses (like Confidentiality, Termination, Governing Law, etc.) in the following text:\n\n`{text}`\n\nIDENTIFIED CLAUSES AND EXPLANATIONS:",
            "refine": "Here is an existing breakdown of contract clauses: `{existing_answer}`. We have a new section of the contract: `{text}`. Add any new clauses and their explanations from this new section to the existing breakdown. Maintain a clear, organized list. If no new clauses are found, return the original breakdown.\n\nUPDATED CLAUSE BREAKDOWN:"
        }
    }
    
    selected_prompts = prompts.get(analysis_type)
    if not selected_prompts:
        return "Invalid analysis type selected."

    initial_prompt_template = ChatPromptTemplate.from_template(selected_prompts["initial"])
    refine_prompt_template = ChatPromptTemplate.from_template(selected_prompts["refine"])

    # Create and run the Refine chain
    try:
        chain = load_summarize_chain(
            llm=llm,
            chain_type="refine",
            question_prompt=initial_prompt_template,
            refine_prompt=refine_prompt_template,
            return_intermediate_steps=False,
            input_key="input_documents",
            output_key="output_text",
        )
        
        result = chain.invoke({"input_documents": docs})
        
        return result.get('output_text', "Analysis failed to produce output.")
    except Exception as e:
        print(f"Error during refine chain execution: {e}")
        return f"An error occurred during analysis. Details: {e}"
    

def get_text_from_file(uploaded_file):
    """
    Extracts text from various file types with a robust, two-stage PDF OCR process.
    Stage 1: Try PyMuPDF for direct text extraction (fast).
    Stage 2: If direct extraction fails, use Pytesseract OCR (slower but handles scanned images).
    """
    try:
        file_bytes = uploaded_file.getvalue()
        file_name_lower = uploaded_file.name.lower() # Convert filename to lowercase for case-insensitive check
        
        if file_name_lower.endswith('.pdf'):
            text_from_pymupdf = ""
            # --- STAGE 1: Try fast, direct text extraction with PyMuPDF ---
            try:
                with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                    for page in doc:
                        text_from_pymupdf += page.get_text()
                
                # Heuristic: If we get a reasonable amount of text, return it.
                if len(text_from_pymupdf.strip()) > 100: 
                    print("Successfully extracted text using PyMuPDF direct method.")
                    return text_from_pymupdf
            except Exception as e:
                print(f"PyMuPDF direct extraction failed: {e}. Falling back to OCR.")

            # --- STAGE 2: Fallback to Pytesseract OCR for scanned PDFs ---
            if not TESSERACT_INSTALLED:
                st.error("This appears to be a scanned PDF, but the Tesseract OCR engine is not installed or not in the system's PATH. Please install Tesseract to process scanned documents.")
                print("Tesseract not found. Cannot perform OCR.")
                return None 

            print("Attempting fallback OCR with Pytesseract...")
            text_from_ocr = ""
            try:
                with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                    for page_num, page in enumerate(doc):
                        pix = page.get_pixmap(dpi=300)  # Render page to an image
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        page_text = pytesseract.image_to_string(img, lang='eng')
                        if page_text:
                            text_from_ocr += page_text + "\n"
                
                if not text_from_ocr.strip():
                     print("Both direct extraction and OCR failed to get text from the PDF.")
                     st.error("Failed to extract text. The PDF may be empty, corrupted, or an unscannable image.")
                     return None

                print("Successfully extracted text using Pytesseract OCR fallback.")
                return text_from_ocr
            except Exception as ocr_error:
                print(f"Pytesseract OCR process failed with an error: {ocr_error}")
                st.error(f"The OCR process failed unexpectedly. Error: {ocr_error}")
                return None
                
        elif file_name_lower.endswith('.docx'):
            document = docx.Document(io.BytesIO(file_bytes))
            return "\n".join([para.text for para in document.paragraphs])
            
        elif file_name_lower.endswith('.txt'):
            return file_bytes.decode("utf-8")
            
        else:
            st.error(f"Unsupported file type: .{uploaded_file.name.split('.')[-1]}")
            return None
            
    except Exception as e:
        st.error(f"A critical error occurred while processing the file: {e}")
        print(f"Unhandled exception in get_text_from_file: {e}")
        return None

def create_pdf_download(title, rich_text):
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        full_html = markdown.markdown(rich_text)
        if "<table>" in full_html:
            parts = full_html.split("<table>")
            before_table_html = parts[0]
            table_and_after = parts[1].split("</table>")
            table_html = table_and_after[0]
            after_table_html = table_and_after[1] if len(table_and_after) > 1 else ""
            if before_table_html:
                pdf.write_html(before_table_html)
            pattern = r"<td.*?>(.*?)</td>\s*<td.*?>(.*?)</td>"
            match = re.search(pattern, table_html, re.DOTALL)
            if match:
                party1_content = match.group(1).strip()
                party2_content = match.group(2).strip()
                party1_text = html2text.html2text(party1_content).strip()
                party2_text = html2text.html2text(party2_content).strip()
                start_y = pdf.get_y()
                pdf.multi_cell(w=pdf.w / 2.2, h=5, txt=party1_text, align='L')
                pdf.set_y(start_y)
                pdf.set_x(pdf.w / 2)
                pdf.multi_cell(w=pdf.w / 2.2, h=5, txt=party2_text, align='L')
            if after_table_html:
                pdf.write_html(after_table_html)
        else:
            pdf.write_html(full_html)
        pdf_bytes = bytes(pdf.output())
        st.download_button(
            label="📄 Download as .PDF",
            data=pdf_bytes,
            file_name=f"{title.replace(' ', '_')}.pdf",
            mime="application/pdf",
            use_container_width=True
        )
    except Exception as e:
        st.error(f"Failed to generate PDF. A formatting error occurred: {e}")

def create_docx_download(title, rich_text):
    text = re.sub(r'##\s*', '', rich_text)
    text = re.sub(r'<.*?>', '', text)
    doc_file = io.BytesIO()
    document = docx.Document()
    document.add_heading(title, 0)
    document.add_paragraph(text)
    document.save(doc_file)
    doc_file.seek(0)
    st.download_button(
        label="📄 Download as .DOCX",
        data=doc_file,
        file_name=f"{title.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

def render_documents_page():
    st.header("Your Saved Documents")
    
    if 'viewing_doc_id' not in st.session_state:
        st.session_state.viewing_doc_id = None

    if st.button("📝 Start a New Document"):
        set_page("Draft")
        st.rerun()

    documents = get_all_documents()
    if not documents:
        st.info("You have no saved documents.")
        return

    for doc in documents:
        doc_id = doc['id']
        is_viewing_this_doc = (st.session_state.viewing_doc_id == doc_id)

        with st.expander(f"{doc['title']} (Last modified: {doc['last_modified'].strftime('%Y-%m-%d %H:%M')})", expanded=is_viewing_this_doc):
            
            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("👁️ View", key=f"view_{doc_id}", use_container_width=True):
                    if is_viewing_this_doc:
                        st.session_state.viewing_doc_id = None
                    else:
                        st.session_state.viewing_doc_id = doc_id
                    st.rerun()
            
            with col2:
                # --- THIS IS THE FIX ---
                # Use the new on_click function that preserves the ID
                if st.button("✍️ Edit", key=f"edit_{doc_id}", use_container_width=True, on_click=edit_document, args=(doc_id,)):
                    # The st.rerun() is now implicitly handled by the on_click callback
                    pass
                # --- END OF FIX ---
            
            with col3:
                if st.button("🗑️ Delete", key=f"del_{doc_id}", type="primary", use_container_width=True):
                    delete_document(doc_id)
                    if st.session_state.viewing_doc_id == doc_id:
                        st.session_state.viewing_doc_id = None
                    st.rerun()

            if is_viewing_this_doc:
                st.markdown("---")
                with st.spinner("Loading document content..."):
                    full_doc = get_document_by_id(doc_id)
                    if full_doc:
                        with st.container(border=True):
                            st.markdown(full_doc['content'])
                    else:
                        st.error("Could not retrieve document content.")
                    
def suggest_revision_for_clause(risky_clause: str):
    """Makes an LLM call to suggest a revision for a specific risky clause."""
    if not llm:
        return "LLM not available. Please check your API key."
        
    system_prompt = "You are an expert legal counsel. Revise a single problematic clause to make it safer, clearer, or more balanced. Provide ONLY the revised clause text, followed by a brief, bulleted explanation of your changes starting with 'Changes Made:'."
    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompt), 
        ("human", "Please revise the following clause:\n\n---\n{clause}\n---")
    ])
    chain = prompt | llm | StrOutputParser()
    return chain.stream({"clause": risky_clause})

def render_draft_page(retriever_instance):
    st.header("Document Drafter & Editor")

    # Session State Management
    if 'draft_page_mode' not in st.session_state: st.session_state.draft_page_mode = 'generate'
    if 'raw_draft_content' not in st.session_state: st.session_state.raw_draft_content = ""
    if 'draft_summary' not in st.session_state: st.session_state.draft_summary = ""
    if 'editable_content' not in st.session_state: st.session_state.editable_content = ""
    if "draft_analysis_result" not in st.session_state: st.session_state.draft_analysis_result = None
    if st.session_state.get("current_doc_id") and ('doc_loaded' not in st.session_state or st.session_state.doc_loaded != st.session_state.current_doc_id):
        doc_data = get_document_by_id(st.session_state.current_doc_id)
        if doc_data: st.session_state.editable_content, st.session_state.editable_title, st.session_state.draft_page_mode, st.session_state.doc_loaded = doc_data['content'], doc_data['title'], 'edit', st.session_state.current_doc_id

    # --- Draft Generation Form ---
    with st.expander("Step 1: Generate a New Draft from Details", expanded=(st.session_state.draft_page_mode == 'generate')):
        with st.form("generator_form"):
            st.subheader("1. Core Document Details")
            col1, col2 = st.columns(2)
            with col1:
                contract_type = st.selectbox("Contract Type", ["Non-Disclosure Agreement", "Service Agreement", "Employment Contract", "Lease Agreement", "Independent Contractor Agreement", "Partnership Agreement", "Software License Agreement"], help="Select the primary type of legal agreement you want to draft.")
                effective_date = st.date_input("Effective Date", value="today", help="The date the contract becomes active.")
                payment_amount = st.text_input("Payment / Consideration", placeholder="e.g., $5,000 USD", help="Specify payment terms or non-monetary consideration.")
            with col2:
                jurisdiction = st.text_input("Governing Jurisdiction", placeholder="e.g., State of California, USA", help="The state or country whose laws will govern the contract.")
                contract_duration = st.text_input("Contract Duration", placeholder="e.g., 1 year from Effective Date")
                termination_notice = st.text_input("Termination Notice Period", placeholder="e.g., 30 days written notice")
            purpose = st.text_area("Purpose / Scope of Agreement", placeholder="e.g., To evaluate a potential business partnership by sharing confidential information. or To define the scope of web development services to be provided...")
            
            st.markdown("---")
            st.subheader("2. The Parties Involved")
            col1, col2 = st.columns(2)
            with col1:
                with st.container(border=True):
                    st.markdown("##### Party 1 Details")
                    party1_name = st.text_input("Name (Company or Individual)", key="p1_name", placeholder="e.g., Innovate Corp")
                    party1_address = st.text_input("Address", key="p1_address", placeholder="e.g., 123 Future Drive, Tech City")
                    party1_incorporation_state = st.text_input("State of Incorporation (if company)", key="p1_inc_state", placeholder="e.g., Delaware")
            with col2:
                with st.container(border=True):
                    st.markdown("##### Party 2 Details")
                    party2_name = st.text_input("Name (Company or Individual)", key="p2_name", placeholder="e.g., Jane Smith")
                    party2_address = st.text_input("Address", key="p2_address", placeholder="e.g., 456 Creative Lane, Design Town")
                    party2_incorporation_state = st.text_input("State of Incorporation (if company)", key="p2_inc_state", placeholder="e.g., California")

            st.markdown("---")
            st.subheader("3. Specific Clauses & Instructions")
            other_details = st.text_area(
                "Add any other specific instructions or details here:", 
                placeholder="e.g., 'The initial payment of $500 is non-refundable.' or 'Specify that all notices must be sent via certified mail.'",
                height=150
            )
            
            st.markdown("---")
            st.subheader("4. Drafting Style & Generation")
            drafting_persona = st.selectbox("Select Drafting Persona", options=["Balanced & Fair", "Pro-Party 1", "Pro-Party 2", "Simple English"])
            
            # --- THIS IS THE CORRECTED PART ---
            if st.form_submit_button("🚀 Generate Draft", type="primary", use_container_width=True):
                if not retriever_instance: 
                    st.error("Cannot generate draft: RAG system not loaded.")
                else:
                    details = {
                        'type': contract_type, 'party1_name': party1_name, 'party2_name': party2_name, 'jurisdiction': jurisdiction,
                        'purpose_of_disclosure': purpose,
                        'effective_date': effective_date.strftime("%B %d, %Y"), 'contract_duration': contract_duration, 'payment_amount': payment_amount,
                        'termination_notice_period': termination_notice, 'other_details': other_details.strip(),
                        'party1_address': party1_address, 'party1_incorporation_state': party1_incorporation_state,
                        'party2_address': party2_address, 'party2_incorporation_state': party2_incorporation_state,
                    }
                    with st.spinner("AI is generating your draft... This may take a moment."):
                        full_draft = st.write_stream(draft_contract_with_rag(details, drafting_persona, retriever_instance))
                    
                    summary_prompt = ChatPromptTemplate.from_template("Provide a brief, one-paragraph executive summary of the following legal contract:\n\n{contract_text}")
                    summary_chain = summary_prompt | llm | StrOutputParser()
                    draft_summary = summary_chain.invoke({"contract_text": full_draft})
                    
                    st.session_state.raw_draft_content = full_draft
                    st.session_state.draft_summary = draft_summary
                    st.session_state.raw_draft_title = f"{contract_type}"
                    st.session_state.draft_page_mode = 'review'
                    st.session_state.doc_loaded, st.session_state.current_doc_id = None, None
                    st.rerun()
            # --- END OF CORRECTION ---

    # --- REVIEW MODE (Unchanged) ---
    if st.session_state.draft_page_mode == 'review':
        # ... (This section is correct and remains the same) ...
        st.subheader("Step 2: Review & Act on Generated Draft")
        st.info("Review the AI-generated text. You can save, download, or load it into the editor for changes.")
        if st.session_state.draft_summary:
            with st.container(border=True):
                st.markdown("**📝 AI Summary of Draft**")
                st.write(st.session_state.draft_summary)
        with st.container(height=400, border=True):
            st.markdown(st.session_state.raw_draft_content)
        st.markdown("---")
        st.subheader("Actions")
        action_col1, action_col2 = st.columns(2)
        with action_col1:
            if st.button("✅ Load into Editor for Changes", type="primary", use_container_width=True):
                st.session_state.editable_content = st.session_state.raw_draft_content
                st.session_state.editable_title = st.session_state.raw_draft_title
                st.session_state.raw_draft_content = ""
                st.session_state.draft_page_mode = 'edit'
                st.rerun()
            if st.button("💾 Save As-Is", use_container_width=True):
                new_id = save_document(st.session_state.raw_draft_title, st.session_state.raw_draft_content)
                st.session_state.current_doc_id = new_id
                st.session_state.draft_page_mode = 'generate'
                st.success(f"Document '{st.session_state.raw_draft_title}' saved successfully!")
                time.sleep(2)
                st.rerun()
        with action_col2:
            create_pdf_download(st.session_state.raw_draft_title, st.session_state.raw_draft_content)
            create_docx_download(st.session_state.raw_draft_title, st.session_state.raw_draft_content)
    # This code goes inside the render_draft_page function

    # --- EDIT MODE ---
    elif st.session_state.draft_page_mode == 'edit':
        st.subheader("Step 3: Edit, Save, and Download")
        
        final_title = st.text_input("Document Title", value=st.session_state.get("editable_title", ""))
        edited_content_html = st_quill(value=st.session_state.get("editable_content", ""), html=True, key="editor")
        edited_content_text = html2text.html2text(edited_content_html)
        
        st.markdown("---")
        st.subheader("Actions")

        # --- THIS IS THE IMPROVED LAYOUT ---
        action_col1, action_col2, action_col3 = st.columns([2, 1, 1])
        
        with action_col1:
            if st.button("💾 Save Document", type="primary", use_container_width=True):
                if final_title and edited_content_text.strip():
                    st.session_state.editable_title = final_title
                    st.session_state.editable_content = edited_content_text
                    if st.session_state.get("current_doc_id"):
                        update_document(st.session_state.current_doc_id, final_title, edited_content_text)
                        st.toast(f"Document '{final_title}' updated!")
                    else:
                        new_id = save_document(final_title, edited_content_text)
                        st.session_state.current_doc_id = new_id
                        st.toast(f"Document '{final_title}' saved!")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.warning("Please provide a title and content before saving.")
        
        with action_col2:
            # The create_..._download functions already use container_width=True
            create_pdf_download(final_title, edited_content_text)
            
        with action_col3:
            create_docx_download(final_title, edited_content_text)

def render_review_page():
    st.header("Summarize & Review an Existing Contract")
    uploaded_file = st.file_uploader("Upload your contract", type=['txt', 'pdf', 'docx'])

    if uploaded_file:
        # The options here are correct (no Risk Assessment)
        analysis_type = st.radio(
            "Select Analysis:",
            ('Chat with Document', 'Executive Summary', 'Clause Breakdown'),
            horizontal=True,
            key=f"analysis_choice_{uploaded_file.name}"
        )

        if analysis_type == 'Chat with Document':
            st.markdown("---")
            st.subheader(f"Chat about: {uploaded_file.name}")
            if "messages" not in st.session_state or st.session_state.get("uploaded_file_name") != uploaded_file.name:
                st.session_state.messages = [AIMessage(content="Hello! I've read the document. How can I help you?")]
                st.session_state.uploaded_file_name = uploaded_file.name
                with st.spinner("Processing document..."):
                    raw_text = get_text_from_file(uploaded_file)
                    if raw_text:
                        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
                        documents = text_splitter.split_text(raw_text)
                        embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
                        vector_store = FAISS.from_texts(documents, embeddings)
                        doc_retriever = vector_store.as_retriever()
                        st.session_state.doc_retriever = doc_retriever
                        st.success("Document processed successfully. You can now ask questions below.")
                    else:
                        st.stop()
            if "doc_retriever" in st.session_state:
                contextualize_q_system_prompt = """Given a chat history and the latest user question which might reference context in the chat history, formulate a standalone question which can be understood without the chat history. Do NOT answer the question, just reformulate it if needed and otherwise return it as is."""
                contextualize_q_prompt = ChatPromptTemplate.from_messages([("system", contextualize_q_system_prompt), MessagesPlaceholder("chat_history"), ("human", "{input}")])
                history_aware_retriever = create_history_aware_retriever(llm, st.session_state.doc_retriever, contextualize_q_prompt)
                
                # --- THIS IS THE CORRECTED PROMPT ---
                # It includes the required {context} placeholder.
                qa_system_prompt = """You are a helpful and professional AI legal assistant.
                    Your user has uploaded a document and is asking questions about its content.
                    Use the following pieces of retrieved context from their document to answer the questions.

                    Your core task is to answer the user's questions based *strictly* on the provided text. 
                    If the answer is not in the context, clearly state that the document does not seem to contain that information. Do not make up information.
                    When answering, it is helpful to quote key phrases from the document to support your answer.

                    Retrieved Context from the User's Document:
                    -----------------
                    {context}
                    -----------------
                    """
                # --- END OF CORRECTION ---

                qa_prompt = ChatPromptTemplate.from_messages([("system", qa_system_prompt), MessagesPlaceholder("chat_history"), ("human", "{input}")])
                
                # This line will now work correctly
                question_answer_chain = create_stuff_documents_chain(llm, qa_prompt)
                
                rag_chain = create_retrieval_chain(history_aware_retriever, question_answer_chain)
                
                for message in st.session_state.messages:
                    with st.chat_message(message.type): st.write(message.content)
                if prompt := st.chat_input("Ask a question about the contract..."):
                    st.session_state.messages.append(HumanMessage(content=prompt))
                    with st.chat_message("human"): st.write(prompt)
                    with st.chat_message("ai"):
                        with st.spinner("Thinking..."):
                            response = rag_chain.invoke({"input": prompt, "chat_history": st.session_state.messages})
                            st.write(response["answer"])
                            st.session_state.messages.append(AIMessage(content=response["answer"]))
        else:
            if st.button("Analyze Document"):
                with st.spinner("Analyzing document... This may take a while for large files."):
                    contract_text = get_text_from_file(uploaded_file)
                    if contract_text:
                        st.session_state.contract_text = contract_text
                        st.session_state.analysis_type = analysis_type
                        st.session_state.analysis_result = analyze_contract(contract_text, analysis_type)

    if 'analysis_result' in st.session_state and st.session_state.get('analysis_type') != 'Chat with Document':
        st.markdown("---")
        current_analysis_type = st.session_state.analysis_type
        result = st.session_state.analysis_result
        st.subheader(f"Analysis Result: {current_analysis_type}")
        st.markdown(result)
        
def clear_review_state():
    keys_to_clear = ['analysis_result', 'analysis_type', 'contract_text', 'messages', 'doc_retriever', 'uploaded_file_name']
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]

def clear_draft_state():
    keys_to_clear = ['draft_page_mode', 'raw_draft_content', 'raw_draft_title', 'editable_content', 'editable_title', 'draft_analysis_result', 'doc_loaded', 'current_doc_id', 'draft_summary']
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]

def set_page(page_name):
    if page_name == "Draft": 
        clear_review_state()
        # Also clear any lingering draft data before starting a new one
        clear_draft_state()
    elif page_name == "Review": 
        clear_draft_state()
    elif page_name == "My Documents": 
        clear_review_state()
        clear_draft_state()
    st.session_state.page = page_name
    
def edit_document(doc_id):
    """Prepares the state to edit a specific document."""
    # First, clear any old state to avoid conflicts
    clear_review_state()
    clear_draft_state()
    
    # Now, set the page and the specific document ID to load
    st.session_state.page = "Draft"
    st.session_state.current_doc_id = doc_id
    
def main_app(retriever, llm):
    st.sidebar.title("AI Legal Assistant")
    st.sidebar.markdown("---")

    # Add Tesseract installation warning to the sidebar
    if not TESSERACT_INSTALLED:
        st.sidebar.warning(
            "Tesseract OCR is not installed or not in PATH. "
            "Text extraction from scanned PDFs will fail. "
            "[See installation instructions](https://github.com/tesseract-ocr/tesseract#installing-tesseract)"
        )
    
    if 'page' not in st.session_state: st.session_state.page = "Home"
    
    def clear_review_state():
        keys_to_clear = ['analysis_result', 'analysis_type', 'contract_text', 'messages', 'doc_retriever', 'uploaded_file_name']
        for key in keys_to_clear:
            if key in st.session_state:
                del st.session_state[key]
    def clear_draft_state():
        keys_to_clear = ['draft_page_mode', 'raw_draft_content', 'raw_draft_title', 'editable_content', 'editable_title', 'draft_analysis_result', 'doc_loaded', 'current_doc_id', 'draft_summary']
        for key in keys_to_clear:
            if key in st.session_state:
                del st.session_state[key]
    def set_page(page_name):
        if page_name == "Draft": clear_review_state()
        elif page_name == "Review": clear_draft_state()
        elif page_name == "My Documents": clear_review_state(); clear_draft_state()
        st.session_state.page = page_name
    st.sidebar.button("🏠 Home", on_click=set_page, args=("Home",), use_container_width=True)
    st.sidebar.button("📄 My Documents", on_click=set_page, args=("My Documents",), use_container_width=True)
    st.sidebar.button("✍️ Draft / Edit Document", on_click=set_page, args=("Draft",), use_container_width=True)
    st.sidebar.button("🔍 Review Contract", on_click=set_page, args=("Review",), use_container_width=True)
    
    # --- CHANGE 3: Add logic to render the home page ---
    if st.session_state.page == "Home":
        render_home_page()
    elif st.session_state.page == "My Documents":
        render_documents_page()
    elif st.session_state.page == "Draft":
        render_draft_page(retriever)
    elif st.session_state.page == "Review":
        render_review_page()

        
if __name__ == "__main__":
    if retriever and llm:
        main_app(retriever, llm)
    else:

        st.error("Application cannot start. RAG models or LLM failed to load. Check the sidebar for errors.")
